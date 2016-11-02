from lxml import html
import requests
from bs4 import BeautifulSoup
from docx import Document

document = Document()
document.add_heading('GECON 2014 Papers', 0)
i = 1
while i < 16:
    r = requests.get("http://link.springer.com/chapter/10.1007/978-3-642-15681-6_" + str(i))
    soup = BeautifulSoup(r.content, "lxml")

    document.add_heading("Paper %s" % i)

    title = soup.find_all("h1", {"class": "ChapterTitle"})
    document.add_heading("Title: ")
    document.add_paragraph(title)

    author_info = soup.find_all('span', {'class': 'AuthorName_container'})
    document.add_heading("Authors:")
    authors = ''
    for info in author_info:
        authors += (info.find('span', {'class': 'AuthorName'}).text) + ', '
        document.add_paragraph(info.find('span', {'class': 'AuthorName'}).text + ' ' + info.find("span", {
            "class": "AuthorsName_affiliation"}).text.replace("Affiliated with", ' '))
        for hrf in info.find_all('a'):
            document.add_paragraph(hrf.get('href'))

    abstract = soup.find("p", {"class": "Para"})
    document.add_heading("Abstract: ")
    document.add_paragraph(abstract.text)

    document.add_heading("Keywords:")
    kwords = ''
    for key in soup.find_all("span", {"class", "Keyword"}):
        kwords += (key.text) + ', '
    document.add_paragraph(kwords[:-2])

    editors = ''
    editors_info = soup.find('ul', {'class': 'editors'}).find_all('a', {'class': 'person'})

    for info in editors_info:
        editors += info.text + ', '

    document.add_heading("To be referenced as:")
    document.add_paragraph(authors[:-2] + '. ' + title + '. ' + "In: Eds.: " + editors[:-2] + '. ' + "GECON, Intl Conf on Grids, Clouds, Systems, and Services. LNCS, vol "
 + soup.find('dd', {'id':'abstract-about-book-series-volume'}).text + ', '
    + soup.find('dd', {'id':'abstract-about-book-chapter-page-ranges'}).text + ', ' + "Springer, 2014.")
    i += 1
document.add_page_break()

document.save('GECON2014.docx')
