from lxml import html
import requests
from bs4 import BeautifulSoup
from docx import Document

edit_r = requests.get("http://link.springer.com/book/10.1007/978-3-319-14609-6")

soup = BeautifulSoup(edit_r.content, "lxml")
editors = ''
editors_info = soup.find('ul', {'class': 'editors'}).find_all('a', {'itemprop': 'name'})
for info in editors_info:
    editors += info.text + ', '


document = Document()
document.add_heading('GECON 2014 Papers', 0)
i = 1
while i < 16:

    r = requests.get("http://link.springer.com/chapter/10.1007/978-3-319-43177-2_" +str(i) + "/fulltext.html")
    soup = BeautifulSoup(r.content, "lxml")

    document.add_heading("Paper %s" % i)

    title = soup.find("h1", {"class": "ChapterTitle"}).text

    author_info = soup.find_all('span', {'class': 'AuthorName_container'})

    authors = ''
    for info in author_info:
        authors += (info.find('span', {'class': 'AuthorName'}).text) + ', '

    document.add_heading("To be referenced as:")
    document.add_paragraph(authors[:-2] + '. ' + title + '. ' + "In: Eds.: " + editors[:-2] + '. ' + "GECON, Intl Conf on Grids, Clouds, Systems, and Services. LNCS, "
                           + str(soup.find('span', {'class': 'vol-info'}).text.replace('of the series Lecture Notes in Computer Science', '')).strip() + ', '
                           + str(soup.find('span', {'class': 'page-numbers-info'}).text).strip() + ', ' + "Springer, 2014.")

    abstract = soup.find("p", {"class": "Para"})
    document.add_heading("Abstract: ")
    document.add_paragraph(abstract.text)

    document.add_heading("Keywords:")
    kwords = ''
    for key in soup.find_all("span", {"class", "Keyword"}):
        kwords += (key.text) + ', '
    document.add_paragraph(kwords[:-2])

    i += 1
document.add_page_break()

document.save('GECON2014.docx')
