from lxml import html
import requests
from bs4 import BeautifulSoup
from docx import Document

# open a file and add the Title
document = Document()
document.add_heading('GECON 2013 Papers', 0)
# url differs only last symbol represented here as i
i = 1
while i < 21:
    print(i)
    # get request from url
    r = requests.get("http://link.springer.com/chapter/10.1007/978-3-319-02414-1_" + str(i))
    soup = BeautifulSoup(r.content, "lxml")

    document.add_heading("Paper %s" % i)

    title = soup.find("h1", {"class": "ChapterTitle"}).text

    author_info = soup.find_all('span', {'class': 'AuthorName_container'})
    authors = ''
    for info in author_info:
        authors += (info.find('span', {'class': 'AuthorName'}).text) + ', '

    editors = ''
    editors_info = soup.find('ul', {'class': 'editors'}).find_all('a', {'class': 'person'})

    for info in editors_info:
        editors += info.text + ', '

    document.add_heading("To be referenced as:")
    document.add_paragraph(authors[:-2] + '. ' + title + '. ' + "In: Eds.: " + editors[:-2] + '. ' + "GECON, Intl Conf on Grids, Clouds, Systems, and Services. LNCS, vol "
                           + soup.find('dd', {'id': 'abstract-about-book-series-volume'}).text + ', '
                           + soup.find('dd', {'id': 'abstract-about-book-chapter-page-ranges'}).text + ', '
                           + "Springer, " + soup.find('dd', {'id': 'abstract-about-book-chapter-copyright-year'}).text + '. ')

    abstract = soup.find("p", {"class": "Para"})
    document.add_heading("Abstract: ")
    document.add_paragraph(abstract.text)

    kwords = ''
    for key in soup.find_all("span", {"class", "Keyword"}):
        kwords += (key.text) + ', '
    if len(kwords)>0:
        document.add_heading("Keywords:")
        document.add_paragraph(kwords[:-2])

    editors = ''
    editors_info = soup.find('ul', {'class': 'editors'}).find_all('a', {'class': 'person'})

    i += 1
document.add_page_break()

document.save('GECON2013.docx')
